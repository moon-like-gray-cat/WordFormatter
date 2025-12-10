from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 打开文档
doc = Document(r"D:\code\WordFormatter\tests\test2.docx")

# ----------------- 第一步：找到第一个一级标题的缩进 -----------------
heading1_indent = None
for p in doc.paragraphs:
    if p.style.name in ("标题 1", "Heading 1"):
        heading1_indent = p.paragraph_format.left_indent or 0
        break

if heading1_indent is None:
    heading1_indent = 0

# ----------------- 第二步：调整所有非标题段落 -----------------
for p in doc.paragraphs:
    if "标题" in p.style.name or "Heading" in p.style.name:
        continue  # 跳过标题

    # 统一左缩进
    p.paragraph_format.left_indent = heading1_indent
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # ----------------- 清理开头多余空格/制表符 -----------------
    if p.text:
        p.text = p.text.lstrip(' \t')  # 去掉开头的空格和制表符

# 保存文档
output_path = r"D:\code\WordFormatter\tests\test2_fixed.docx"
doc.save(output_path)
print(f"文档缩进已统一并清理开头空格，保存为 {output_path}")
