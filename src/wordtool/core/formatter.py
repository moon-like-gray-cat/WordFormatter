from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
import win32com.client as win32
import re
from docx.enum.style import WD_STYLE_TYPE



TITLE_FORMATS = [
    "一", "一、", "（一）", "（一）、", "（一）.",
    "（1）", "（1）、", "（1）.", "1", "1.", "1、","1.1","1.1.1",
    "a", "a.", "A", "A.", "①", "I", "I.", "（I）"
]

_FORMAT_TO_REGEX = {
    # ------------------ 中文大写数字类 -------------------
    "一": r"^[一二三四五六七八九十]+\s*",  # 匹配 "一 " 或 "二 " (无标点)
    "一、": r"^[一二三四五六七八九十]+[、\.]\s*",  # 匹配 "一、", "二.", "三 " (带顿号或点号)
    "（一）": r"^（[一二三四五六七八九十]+）\s*",  # 匹配 "(一) ", "(二) "

    # ------------------ 阿拉伯数字类 ----------------------
    "1": r"^\d+\s*",  # 匹配 "1 " 或 "2 " (无标点)    "1": r"^\d+(?![\.])\s*",

    "1.": r"^\d+\.(?!\d)\s*",

    "1.1": r"^\d+\.\d+(?!\.)\s*",

    "1.1.1": r"^\d+\.\d+\.\d+(?!\.)\s*",

    "1.1.1.1": r"^\d+\.\d+\.\d+\.\d+(?!\.)\s*",

    "（1）": r"^（\d+）\s*",  # 匹配 "(1) ", "(2) "

    # ------------------ 字母和罗马数字类 --------------------
    "a": r"^[a-z]{1,2}\s*",  # 匹配 "a " 或 "b "
    "a.": r"^[a-z]{1,2}[、\.]\s*",  # 匹配 "a.", "b、"
    "A": r"^[A-Z]{1,2}\s*",  # 匹配 "A " 或 "B "
    "A.": r"^[A-Z]{1,2}[、\.]\s*",  # 匹配 "A.", "B、"
    "I": r"^[IVXLCDM]+\s*",  # 匹配 "I " 或 "II "
    "I.": r"^[IVXLCDM]+[、\.]\s*",  # 匹配 "I.", "II、"
    "（I）": r"^（[IVXLCDM]+）\s*",  # 匹配 "(I) ", "(II) "

    # ------------------ 特殊符号类 -------------------------
    "①": r"^[①②③④⑤⑥⑦⑧⑨⑩]+\s*",  # 匹配带圈数字
}


def extract_pt(size_str: str) -> float:
    """
    将 config 里的 "四号 (14pt)" / "五号 (10.5pt)" 转成数字 pt。
    如果没有括号，尝试强行解析数字。
    """
    m = re.search(r"\(([\d.]+)pt\)", size_str)
    if m:
        return float(m.group(1))

    # fallback 强行取前面的数字
    m = re.search(r"([\d.]+)", size_str)
    if m:
        return float(m.group(1))

    return 12.0

def ensure_outline_level(paragraph, level: int):
    """
    仅设置 outline level（用于目录 / 导航窗格）
    不使用 Heading 样式
    """

    # level=1 → outlineLvl=0
    outline_val = level - 1
    if outline_val < 0:
        return

    pPr = paragraph._element.get_or_add_pPr()

    # 防止重复 outlineLvl
    existing = pPr.find(qn("w:outlineLvl"))
    if existing is not None:
        pPr.remove(existing)

    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), str(outline_val))
    pPr.append(outline)

def paragraph_has_drawing(paragraph):
    """
    判断段落中是否包含嵌入型图片（inline image / drawing）
    """
    return bool(
        paragraph._element.xpath(".//w:drawing")
    )

class WordFormatter:
    def __init__(self, file_path, config: dict):
        self.file_path = file_path
        self.config = config
        self.titles = config.get("titles", {})
        self.body = config.get("body", {})
        self.figure = config.get("figure", {})
        self.table = config.get("table", {})
        #读取是否展开自动编号（默认 True）
        self.expand_numbering = (
            config.get("options", {})
            .get("expand_numbering", True)
        )


    def _expand_numbering(self, input_path, output_path):
        """
        使用 Word COM 将自动编号转为真实文本
        - 若系统无 Office / win32com 不可用 → 自动跳过
        - 若路径非法 / 打不开 → 自动跳过
        - 任何异常不影响主流程
        """
        import os

        # 路径标准化（Word COM 对路径极其敏感）
        input_path = os.path.abspath(input_path)
        output_path = os.path.abspath(output_path)

        # 输入文件不存在，直接跳过
        if not os.path.exists(input_path):
            print(f"[WordFormatter] Skip expand numbering: file not found: {input_path}")
            return input_path

        # 尝试导入 win32com（判断是否有 Office 能力）
        try:
            import win32com.client as win32
        except ImportError:
            print("[WordFormatter] Skip expand numbering: win32com not available (no Office?)")
            return input_path

        word = None
        doc = None

        try:
            # 使用 DispatchEx，避免抢占已有 Word 实例
            word = win32.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0  # 禁止弹窗

            # 打开文档（只读 = False）
            doc = word.Documents.Open(input_path, ReadOnly=False)

            # 展开自动编号
            doc.ConvertNumbersToText()

            # 保存为新文件
            doc.SaveAs(output_path)

            return output_path

        except Exception as e:
            print(f"[WordFormatter] Error expanding numbering: {e}")
            return input_path

        finally:
            # 资源清理（必须）
            try:
                if doc is not None:
                    doc.Close(False)
            except Exception:
                pass

            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass

    # ----------------------------------------------------------------------
    # 设置文本 run 样式（图片 run 跳过）
    # ----------------------------------------------------------------------

    def _set_run_style(self, run, cn_font, cn_size, bold, en_font=None, en_size=None):
        # 1. 跳过图片 run
        if run._element.xpath(".//w:drawing"):
            return

        # 获取或创建 rPr (Run Properties) 元素
        rPr = run._element.get_or_add_rPr()

        # --- 2. 字体处理 (中西文分离) ---
        rFonts = rPr.get_or_add_rFonts()
        if cn_font:
            # 设置东亚字体（中文）
            rFonts.set(qn("w:eastAsia"), cn_font)
        if en_font:
            # 设置西文字体（ASCII 字符范围）
            rFonts.set(qn("w:ascii"), en_font)
            # 设置高 ANSI 字体（通常也要指向西文字体，确保符号一致）
            rFonts.set(qn("w:hAnsi"), en_font)

        # --- 3. 字号处理 (中西文分离) ---
        # Word XML 中字号单位是半磅 (Half-points)，即 Pt * 2
        def set_xml_size(rPr_elem, size_pt, tag_name):
            existing = rPr_elem.find(qn(tag_name))
            if existing is not None:
                rPr_elem.remove(existing)
            new_tag = OxmlElement(tag_name)
            new_tag.set(qn("w:val"), str(int(size_pt * 2)))
            rPr_elem.append(new_tag)

        if en_size:
            # w:sz 对应西文字号
            set_xml_size(rPr, en_size, "w:sz")
        if cn_size:
            # w:szCs 对应复杂字符/东亚字符字号
            set_xml_size(rPr, cn_size, "w:szCs")

        # --- 4. 其他样式 (加粗、颜色) ---
        run.font.bold = bool(bold)
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _clean_numbering_spaces(self, doc):
        for para in doc.paragraphs:

            if not para.runs:
                continue

            # 找到第一个【真正有文本的 run】
            first_text_run = None
            for run in para.runs:
                if run.text and run.text.strip():
                    first_text_run = run
                    break

            if not first_text_run:
                continue

            original = first_text_run.text.lstrip(" \t")
            new_text = original

            for i in range(4, 0, -1):
                key = f"title{i}"
                fmt = self.titles.get(key, {}).get("format", "")
                if not fmt:
                    continue

                regex = _FORMAT_TO_REGEX.get(fmt)
                if regex:
                    new_text = re.sub(f"({regex})\\s+", r"\1", original)
                    break

            if new_text != original:
                first_text_run.text = new_text

    # ----------------------------------------------------------------------
    # 标题层级检测
    # ----------------------------------------------------------------------
    def _detect_level(self, text):
        # 1. 预处理文本：标准化括号和去除隐藏字符
        normalized_text = self._normalize_brackets(text.strip())
        normalized_text = re.sub(r'^[\s\x00-\x1f]+', '', normalized_text)
        if not normalized_text:
            return 0

        for i in range(4, 0, -1):
            key = f"title{i}"
            # 2. 从 JSON 配置中获取用户设定的标识 (例如 "（1）" 或 "1.")
            format_key = self.titles.get(key, {}).get("format", "")

            # 3. 查表获取对应的正则表达式字符串
            # 注意：这里我们使用全局的 _FORMAT_TO_REGEX 字典
            regex_pattern = _FORMAT_TO_REGEX.get(format_key)

            if regex_pattern:
                # 4. 使用 re.match 进行匹配（执行前缀匹配）
                try:
                    if re.match(regex_pattern, normalized_text):
                        return i
                except re.error as e:
                    # 提示配置中正则错误
                    print(f"Warning: Invalid regex pattern used for {key}: {e}")
                    continue
        return 0

    # ----------------------------------------------------------------------
    # 获取样式
    # ----------------------------------------------------------------------
    def _get_style(self, level):
        """
        标题样式 = body 默认值 + title 覆盖
        """
        if level == 0:
            return self.body

        key = f"title{level}"
        title_cfg = self.titles.get(key, {})

        merged = dict(self.body)  # 继承正文行距
        merged.update(title_cfg)  # 标题只覆盖字体/字号等

        return merged



    # ----------------------------------------------------------------------
    # 应用样式到段落
    # ----------------------------------------------------------------------
    def _apply_style(self, paragraph, level, doc, caption_type=None):
        """
        paragraph: 段落对象
        level: 0=正文，1~n=标题
        caption_type: "caption" 表示图/表标题
        """

        # ------------------------------------------------------------------
        # 1. 获取样式配置
        # ------------------------------------------------------------------
        if level == 0 and caption_type == "caption":
            style_cfg = self.config.get("caption", {})
        else:
            style_cfg = self._get_style(level)

        # ------------------------------------------------------------------
        # 2. 统一使用 Normal，彻底摆脱 Heading
        # ------------------------------------------------------------------
        paragraph.style = doc.styles["Normal"]

        # ------------------------------------------------------------------
        # 3. 标题：写入 outline level（给目录识别）
        # ------------------------------------------------------------------
        if level > 0:
            ensure_outline_level(paragraph, level)

        # ------------------------------------------------------------------
        # 4. Run 级别：字体 / 字号 / 加粗
        # ------------------------------------------------------------------
        # 1. 统一字体字号配置获取
        cn_font = style_cfg.get("font", "宋体")
        cn_size = extract_pt(style_cfg.get("size", "12pt"))
        bold = bool(style_cfg.get("bold", False))

        en_cfg = self.config.get("en_font", {})
        en_font = en_cfg.get("font", "Times New Roman")

        if level == 0 and caption_type != "caption":
            en_size = extract_pt(en_cfg.get("size", "12pt"))
        else:
            en_size = cn_size  # 标题及图表说明，西文随中文

        for run in paragraph.runs:
            self._set_run_style(
                run,
                cn_font=cn_font,
                cn_size=cn_size,
                bold=bold,
                en_font=en_font,
                en_size=cn_size
            )
        # ------------------------------------------------------------------
        # 5. 段落级别：行距（完全 JSON 驱动）
        # ------------------------------------------------------------------
        pf = paragraph.paragraph_format
        pPr = paragraph._element.get_or_add_pPr()

        # 清除可能残留的 spacing（包括复制粘贴遗留）
        spacing = pPr.find(qn("w:spacing"))
        if spacing is not None:
            pPr.remove(spacing)

        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        line_rule = style_cfg.get("line_rule", "多倍行距")
        spacing_val = float(style_cfg.get("spacing", 1.25))

        if line_rule == "多倍行距":
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = spacing_val
        else:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(spacing_val)

        # ------------------------------------------------------------------
        # 6. 对齐 & 缩进规则
        # ------------------------------------------------------------------
        if level > 0:
            # 标题：全部顶格
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(0)
            pf.hanging_indent = Pt(0)

        elif caption_type == "caption":
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pf.first_line_indent = Pt(0)

        else:
            # 正文缩进留给 _normalize_paragraph_indent 统一处理
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # ----------------------------------------------------------------------
    # 将英文括号转中文括号
    # ----------------------------------------------------------------------
    def _normalize_brackets(self, text):
        text = text.replace("(", "（").replace(")", "）")
        return text

    # ----------------------------------------------------------------------
    # 处理图题和表题
    # ----------------------------------------------------------------------
    def _preprocess_captions(self, doc):
        """
        处理已有图题和表题：
        - 图片下方图题
        - 表格上方表题
        """
        paragraphs = doc.paragraphs
        for i, para in enumerate(paragraphs):
            # 图片下方图题
            if para._element.xpath(".//w:drawing"):
                if i + 1 < len(paragraphs):
                    next_text = paragraphs[i + 1].text.strip()
                    # 以"图"开头，后面跟着数字
                    if re.match(r'^图\s*\d+.*', next_text):
                        caption_para = paragraphs[i + 1]
                    self._apply_style(caption_para, level=0,doc=doc,caption_type="caption")
                    caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 表格上方表题
            next_elem = para._element.getnext()
            if next_elem is not None and next_elem.tag.endswith("tbl"):
                if re.match(r'^表\s*\d+.*', para.text.strip()):
                    caption_para = para
                    self._apply_style(caption_para, level=0,doc=doc, caption_type="caption")
                    caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def adjust_line_spacing_for_images(self,doc):
        """
        若段落中包含图片，且行距为固定值（Exactly）并且较小，
        则自动切换为单倍行距，避免图片被裁剪。
        """

        for para in doc.paragraphs:

            if not paragraph_has_drawing(para):
                continue

            fmt = para.paragraph_format

            # 只处理固定行距
            if fmt.line_spacing_rule != WD_LINE_SPACING.EXACTLY:
                continue

            # line_spacing 可能为 None（异常文档）
            if not fmt.line_spacing:
                continue

            # python-docx 中 line_spacing 为 Length（EMU）
            try:
                spacing_pt = fmt.line_spacing.pt
            except Exception:
                continue

            fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
            fmt.line_spacing = None

    def _normalize_paragraph_indent(self, doc):
        body_size_str = self.body.get("size", "小四 (12pt)")
        body_pt = extract_pt(body_size_str)

        # 定义 two_char_indent：2个汉字的宽度 = 字号 * 2
        two_char_indent = Pt(body_pt * 2)

        for p in doc.paragraphs:
            fmt = p.paragraph_format
            # 1. 尝试通过正则检测
            level_by_regex = self._detect_level(p.text)
            # 2. 尝试从 XML 属性中直接获取大纲级别 (0 代表 1 级标题)
            pPr = p._element.find(qn("w:pPr"))
            has_outline = False
            if pPr is not None:
                if pPr.find(qn("w:outlineLvl")) is not None:
                    has_outline = True

            # 只要满足其一，就认定是标题，强制顶格
            if level_by_regex > 0 or has_outline:
                fmt = p.paragraph_format
                fmt.first_line_indent = Pt(0)
                fmt.left_indent = Pt(0)
                fmt.hanging_indent = Pt(0)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                continue

            fmt.left_indent = Pt(0)
            fmt.first_line_indent = two_char_indent
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # ----------------------------------------------------------------------
    # 保存文档
    # ----------------------------------------------------------------------

    def save(self, output_path):
        try:

            # ----------------- 1. 是否展开自动编号 -----------------
            if self.expand_numbering:
                expanded_path = self._expand_numbering(
                    self.file_path,
                    output_path.replace(".docx", "_expanded.docx")
                )
            else:
                expanded_path = self.file_path
            # ----------------- 2. 用 python-docx 打开展开后的文档 -----------------
            doc = Document(expanded_path)

            # 清理编号和标题间多余空格
            self._clean_numbering_spaces(doc)


            # ----------------- 3. 标题括号规范化 -----------------
            for para in doc.paragraphs:
                for run in para.runs:
                    if not run._element.xpath(".//w:drawing"): # 检查run是否包含图片
                        run.text = self._normalize_brackets(run.text) # 对纯文本进行处理


            # ----------------- 4. 应用样式（标题/正文） -----------------
            for para in doc.paragraphs:
                level = self._detect_level(para.text)
                self._apply_style(para, level, doc)

            # 修复图片 + 固定行距冲突
            self.adjust_line_spacing_for_images(doc)
            # 段落缩进，以及 两端对齐
            self._normalize_paragraph_indent(doc)
            # ----------------- 5. 处理已有图题/表题 -----------------
            self._preprocess_captions(doc)

            # ----------------- 6. 保存最终文档 -----------------
            doc.save(output_path)
            print(f"文档保存成功：{output_path}")
            return True

        except Exception as e:
            print(f"Error saving document: {e}")
            return False